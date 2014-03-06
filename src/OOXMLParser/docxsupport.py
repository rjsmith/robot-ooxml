# Copyright (c) 2014 Richard Smith, https://github.com/rjsmith
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#    http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from robot.errors import DataError

try:
    from docx import Document
except ImportError:
    raise DataError("Using docx test data requires having "
                    "'python-docx' module version 0.3 or newer installed.")

class DocxDocumentReader:
    '''
    Wrapper class around python-docx API calls to inspect structure of docx file
    '''
    def __init__(self, docxfile):
        self.document = Document(docxfile)
        
    def getTables(self):
        return self.document.tables
    
    def getTableRows(self, table):
        return table.rows
        
    def getTextOfRowCells(self, row):
        cellTexts = []
        for cell in row.cells:
            paragraphs = cell.paragraphs
            cellText = self._concatenateParagraphs(paragraphs)  
            cellTexts.append(cellText)
        return cellTexts
    
    def _concatenateParagraphs(self, paragraphs):
        """
        Concatenate all text in list of paragraphs
        """
        text = ''
        for paragraph in paragraphs:
            text += paragraph.text
        return text
        
        
        